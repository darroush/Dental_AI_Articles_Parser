import os
from openai import OpenAI
import pandas as pd
from docx import Document
from docx.shared import Inches
from docx.enum.section import WD_ORIENT
from tqdm import tqdm
from helper import *

# Initialize OpenAI API key from environment variable
client = OpenAI(
  api_key= os.getenv("OPENAI_API_KEY")
)

# TODO:  
Entrez.email = "mostafadesoki86@gmail.com"
time_interval = '2024/10/15:2024/10/31'
search_query = f"""((("Artificial Intelligence"[Mesh] OR "Artificial Intelligence"[tw] OR "Machine Learning"[tw] OR "Convolutional neural network*"[tw] OR "Deep Learning"[tw] AND ({time_interval}[pdat])) AND 
                ("Dentistry"[Mesh] OR Dentistry[tw] OR Dental[tw] AND ({time_interval}[pdat])) NOT "systematic review"[Filter]) NOT "review"[Filter])"""
gpt_model = "gpt-4o-mini"

# Fetch the articles from PubMed database and export as Excel file
articles = fetch_pubmed_data(search_query, max_records=100)
pubmed_parser(articles)

# Load the Excel file
file_path = 'parser_output.xlsx'
df = pd.read_excel(file_path)

# Define function to process each abstract with OpenAI API
def process_abstract(abstract):
    prompt = f"""
    Here is an abstract of a research paper:
    {abstract}
    
    Please extract the following information:
    - Objectives: (the aim of the study)
    - Study Type: (classification, object detection, semantic segmentation, etc.)
    - Data Type: (describe the data used, e.g., "1300 panoramic radiographs of patients")
    - Model Architecture: (e.g., "CNN (Yolo V8)", "SVM", etc.)
    - Metrics: (main performance indicators like IoU or F1 score)
    - Conclusion: (main outcome of the study)
    
    If a field is not mentioned in the abstract, respond with "N/A".
    """

    try:
        response = client.chat.completions.create(
            model= gpt_model,
            messages=[{"role": "system", "content": "You are a helpful assistant for processing research abstracts."},
                      {"role": "user", "content": prompt}],
            max_tokens=300,
            temperature=0
        )
        
        # Corrected access to response content
        content = response.choices[0].message.content.strip()
        # print("Raw response content:\n", content)  # Debug line to see the exact response format

                
        # Parse the response content into a dictionary
        sections = {}
        for line in content.split("\n"):
            # Remove leading "- " if it exists, and then split by ":"
            line = line.lstrip("- ").strip()
            if ":" in line:
                key, value = line.split(":", 1)
                sections[key.strip()] = value.strip()
        return sections    
    except Exception as e:
        print(f"Error processing abstract: {e}")
        return None

# Loop through each row and process abstracts
processed_data = []

for index, row in tqdm(df.iterrows(), total=len(df)):
    title = row.get('Title', 'N/A')
    author = row.get('Author', 'N/A')
    journal = row.get('Journal', 'N/A')
    DOI = row.get('DOI', 'N/A')
    abstract = row.get('Abstract', None)

    if abstract:
        extracted_info = process_abstract(abstract) or {}
        
        # Append to list
        processed_data.append({
            'Index': index + 1,
            'Title': title,
            'Author': author,
            'Journal': journal,
            'DOI': DOI,
            'Objectives': extracted_info.get("Objectives", "N/A"),
            'Study Type': extracted_info.get("Study Type", "N/A"),
            'Data Type': extracted_info.get("Data Type", "N/A"),
            'Model Architecture': extracted_info.get("Model Architecture", "N/A"),
            'Metrics': extracted_info.get("Metrics", "N/A"),
            'Conclusion': extracted_info.get("Conclusion", "N/A")
        })

# Convert list to DataFrame and save to Word document
processed_df = pd.DataFrame(processed_data)
doc = Document()
section = doc.sections[0]

# Set orientation to landscape for A3 paper size
section.orientation = WD_ORIENT.LANDSCAPE
section.page_height = Inches(11.7)  # A3 height in inches
section.page_width = Inches(16.5)   # A3 width in inches

doc.add_heading('Extracted Research Paper Information', level=1)

# Define column headers for Word table
column_headers = ['Index', 'Title', 'Author', 'Journal', 'DOI', 'Objectives', 'Study Type', 'Data Type', 'Model Architecture', 'Metrics', 'Conclusion']
column_widths = [Inches(0.56), Inches(1.16), Inches(1.16), Inches(1.18), Inches(0.6), Inches(2), Inches(0.88), Inches(1.6), Inches(1.4), Inches(1.3), Inches(2)]

# Create table and set header row with specified column widths
table = doc.add_table(rows=1, cols=len(column_headers))
hdr_cells = table.rows[0].cells
for i, header in enumerate(column_headers):
    hdr_cells[i].text = header

# Fill in table rows with specified column widths
for i, row in processed_df.iterrows():
    row_cells = table.add_row().cells
    for j, header in enumerate(column_headers):
        row_cells[j].text = str(row[header])

# Set column widths using the table object directly
for i, width in enumerate(column_widths):
    for row in table.rows:
        row.cells[i].width = width

table.autofit = False  # Disable auto fit to allow manual width setting

# Save document
time_interval_dashes = time_interval.replace("/", "-").replace(":", "-to-")
output_path = f'DentalAIarticles_summary_{time_interval_dashes}_{gpt_model}.docx'
doc.save(output_path)
print(f"Document saved as {output_path}")
